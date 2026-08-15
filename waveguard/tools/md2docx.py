#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Minimal Markdown -> DOCX converter for 护院鹅 docs.
Handles: headings, blockquotes, tables, fenced code, bullet/checkbox lists,
inline **bold** and `code`. Sets Chinese-friendly fonts.
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1]
DST = sys.argv[2]

CJK_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"

def set_cjk(run, font=CJK_FONT):
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)

def add_runs(paragraph, text, base_font=CJK_FONT):
    """Parse inline **bold** and `code`, add runs."""
    # split by bold and code while keeping delimiters
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2]); run.bold = True; set_cjk(run, base_font)
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1]); set_cjk(run, MONO_FONT); run.font.size = Pt(10)
        else:
            run = paragraph.add_run(tok); set_cjk(run, base_font)

def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells

doc = Document()
# default style font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

i = 0
n = len(lines)
in_code = False
code_buf = []
while i < n:
    line = lines[i]
    # fenced code
    if line.strip().startswith("```"):
        if not in_code:
            in_code = True; code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run("\n".join(code_buf))
            set_cjk(run, MONO_FONT); run.font.size = Pt(9.5)
        i += 1; continue
    if in_code:
        code_buf.append(line); i += 1; continue

    # table block
    if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-\|]+\|\s*$", lines[i+1]):
        header = parse_table_row(line)
        i += 2  # skip separator
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append(parse_table_row(lines[i])); i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], h)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for r in rows:
            cells = table.add_row().cells
            for j in range(len(header)):
                val = r[j] if j < len(r) else ""
                cells[j].paragraphs[0].text = ""
                add_runs(cells[j].paragraphs[0], val)
        doc.add_paragraph()
        continue

    stripped = line.strip()
    # headings
    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1)); text = m.group(2)
        h = doc.add_heading(level=min(level, 4))
        h.text = ""
        add_runs(h, text)
        i += 1; continue
    # horizontal rule
    if stripped in ("---", "***", "___"):
        i += 1; continue
    # blockquote
    if stripped.startswith(">"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        add_runs(p, stripped.lstrip(">").strip())
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1; continue
    # checkbox / bullet list
    mb = re.match(r"^[-*]\s+(\[[ xX]\]\s+)?(.*)$", stripped)
    if mb:
        p = doc.add_paragraph(style="List Bullet")
        prefix = ""
        if mb.group(1):
            checked = "x" in mb.group(1).lower()
            prefix = "☑ " if checked else "☐ "
        add_runs(p, prefix + mb.group(2))
        i += 1; continue
    # numbered list
    mn = re.match(r"^\d+\.\s+(.*)$", stripped)
    if mn:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, mn.group(1))
        i += 1; continue
    # blank
    if stripped == "":
        i += 1; continue
    # normal paragraph
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(DST)
print("SAVED:", DST)
